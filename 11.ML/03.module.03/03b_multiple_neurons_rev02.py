import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import torch
import torch.nn as nn
from torch import sigmoid
import numpy as np

def main():
    torch.manual_seed(123)
    class Net(nn.Module):
        def __init__(self, D_in, H, D_out):
            super(Net, self).__init__()
            self.linear1 = nn.Linear(D_in, H)
            self.linear2 = nn.Linear(H, D_out)
            self.a1 = None
            self.l1 = None
            self.l2 = None
        def forward(self, x):
            self.l1 = self.linear1(x)
            self.a1 = sigmoid(self.l1)
            self.l2 = self.linear2(self.a1)
            yhat = sigmoid(self.linear2(self.a1))
            return yhat
    def train(Y, X, model, optimizer, criterion, epochs=600):
        cost = []
        for epoch in range(epochs):
            total = 0
            for y, x in zip(Y, X):
                yhat = model(x)
                loss = criterion(yhat, y)
                loss.backward()
                optimizer.step()
                optimizer.zero_grad()
                total += loss.item()
            cost.append(total)
            if epoch % 200 == 0:
                plt.plot(X.numpy(), model(X).detach().numpy(), label=('epoch ' + str(epoch)))
                plt.plot(X.numpy(), Y.numpy(), 'r')
                plt.xlabel('x')
                plt.legend()
                plt.show()
        return cost
    # Different data: more samples, different range
    X = torch.arange(-10, 10, 0.5).view(-1, 1).type(torch.FloatTensor)
    Y = torch.zeros(X.shape[0])
    Y[(X[:, 0] > -2) & (X[:, 0] < 2)] = 1.0
    D_in = 1
    H = 5  # more hidden neurons
    D_out = 1
    learning_rate = 0.01
    model = Net(D_in, H, D_out)
    optimizer = torch.optim.Adam(model.parameters(), lr=learning_rate)
    def criterion_cross(outputs, labels):
        out = -1 * torch.mean(labels * torch.log(outputs) + (1 - labels) * torch.log(1 - outputs))
        return out
    cost_cross = train(Y, X, model, optimizer, criterion_cross, epochs=600)
    plt.plot(cost_cross)
    plt.xlabel('epoch')
    plt.title('cross entropy loss')
    plt.show()
    x = torch.tensor([1.0])
    yhat = model(x)
    print('Prediction for x=1.0:', yhat)
    X_ = torch.tensor([[1.0], [3.0], [-3.0]])
    Yhat = model(X_)
    print('Predictions for X_:', Yhat)
    Yhat = Yhat > 0.5
    print('Thresholded predictions:', Yhat)

output_buffer = io.StringIO()
plots = []
_original_show = plt.show
def save_and_collect(*args, **kwargs):
    fname = f"{os.path.splitext(os.path.basename(__file__))[0]}_plot_{len(plots)+1}.png"
    plt.savefig(fname)
    plots.append(fname)
    plt.close()
plt.show = save_and_collect
doc = Document()
doc.add_heading('Script Output', 0)
with redirect_stdout(output_buffer):
    main()
plt.show = _original_show
output_text = output_buffer.getvalue()
doc.add_paragraph(output_text)
for plot_file in plots:
    doc.add_picture(plot_file, width=Inches(5))
    os.remove(plot_file)
docx_name = f"{os.path.splitext(os.path.basename(__file__))[0]}.docx"
doc.save(docx_name) 