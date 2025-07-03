# <h1>Simple One Hidden Layer Neural Network</h1>
# <h2>Objective</h2><ul><li> How to create simple Neural Network in pytorch.</li></ul> 
# <h2>Table of Contents</h2>
# <p>In this lab, you will use a single-layer neural network to classify non linearly seprable data in 1-Ddatabase.</p>
# <ul>
#     <li><a href="#Model">Neural Network Module and Training Function</a></li>
#     <li><a href="#Makeup_Data">Make Some Data</a></li>
#     <li><a href="#Train">Define the Neural Network, Criterion Function, Optimizer, and Train the Model</a></li>
# </ul>
# <p>Estimated Time Needed: <strong>25 min</strong></p>
# <hr>
# <h2>Preparation</h2>
# We'll need the following libraries

import torch
import torch.nn as nn
from torch import sigmoid
import matplotlib.pylab as plt
import numpy as np
import os
from docx import Document
from docx.shared import Inches

torch.manual_seed(0)

# Set device
device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')

# Used for plotting the model
def PlotStuff(X, Y, model, epoch, leg=True, save_path=None):
    plt.figure()
    plt.plot(X.cpu().numpy(), model(X).detach().cpu().numpy(), label=('epoch ' + str(epoch)))
    plt.plot(X.cpu().numpy(), Y.cpu().numpy(), 'r')
    plt.xlabel('x')
    if leg:
        plt.legend()
    if save_path:
        plt.savefig(save_path)
    plt.close()

# Define the class Net
class Net(nn.Module):
    def __init__(self, D_in, H, D_out):
        super(Net, self).__init__()
        self.linear1 = nn.Linear(D_in, H)
        self.linear2 = nn.Linear(H, D_out)
        self.a1 = None
        self.l1 = None
        self.l2 = None
    def forward(self, x):
        self.l1 = self.linear1(x)
        self.a1 = sigmoid(self.l1)
        self.l2 = self.linear2(self.a1)
        yhat = sigmoid(self.linear2(self.a1))
        return yhat

def train(Y, X, model, optimizer, criterion, epochs=1000, doc=None, fig_dir="figs"):
    cost = []
    total = 0
    os.makedirs(fig_dir, exist_ok=True)
    for epoch in range(epochs):
        total = 0
        for y, x in zip(Y, X):
            yhat = model(x)
            loss = criterion(yhat, y)
            loss.backward()
            optimizer.step()
            optimizer.zero_grad()
            total += loss.item()
        cost.append(total)
        if epoch % 300 == 0:
            fig_path = os.path.join(fig_dir, f"model_epoch_{epoch}.png")
            PlotStuff(X, Y, model, epoch, leg=True, save_path=fig_path)
            if doc:
                doc.add_paragraph(f"Model output at epoch {epoch}")
                doc.add_picture(fig_path, width=Inches(4))
            model(X)
            plt.figure()
            plt.scatter(model.a1.detach().cpu().numpy()[:, 0], model.a1.detach().cpu().numpy()[:, 1], c=Y.cpu().numpy().reshape(-1))
            plt.title('activations')
            act_fig_path = os.path.join(fig_dir, f"activations_epoch_{epoch}.png")
            plt.savefig(act_fig_path)
            plt.close()
            if doc:
                doc.add_paragraph(f"Activations at epoch {epoch}")
                doc.add_picture(act_fig_path, width=Inches(4))
    return cost

# Make some data
X = torch.arange(-20, 20, 1).view(-1, 1).type(torch.FloatTensor).to(device)
Y = torch.zeros(X.shape[0]).to(device)
Y[(X[:, 0] > -4) & (X[:, 0] < 4)] = 1.0

def criterion_cross(outputs, labels):
    out = -1 * torch.mean(labels * torch.log(outputs) + (1 - labels) * torch.log(1 - outputs))
    return out

# Set up docx document
doc = Document()
doc.add_heading('Simple One Hidden Layer Neural Network - GPU Version', 0)
doc.add_paragraph(f"Device used: {device}")

# Train the model
D_in = 1
H = 2
D_out = 1
learning_rate = 0.1
model = Net(D_in, H, D_out).to(device)
optimizer = torch.optim.SGD(model.parameters(), lr=learning_rate)
doc.add_heading('Training with Cross Entropy Loss', level=1)
cost_cross = train(Y, X, model, optimizer, criterion_cross, epochs=1000, doc=doc)

# Plot and save the loss curve
plt.figure()
plt.plot(cost_cross)
plt.xlabel('epoch')
plt.title('cross entropy loss')
loss_fig_path = "figs/loss_curve.png"
plt.savefig(loss_fig_path)
plt.close()
doc.add_paragraph("Loss curve:")
doc.add_picture(loss_fig_path, width=Inches(4))

# Predictions
x = torch.tensor([0.0]).to(device)
yhat = model(x)
doc.add_paragraph(f'Prediction for x=0.0: {yhat.detach().cpu().numpy()}')

X_ = torch.tensor([[0.0], [2.0], [3.0]]).to(device)
Yhat = model(X_)
doc.add_paragraph(f'Predictions for X_: {Yhat.detach().cpu().numpy()}')

Yhat_thresh = Yhat > 0.5
doc.add_paragraph(f'Thresholded predictions: {Yhat_thresh.detach().cpu().numpy()}')

doc.save("03a_simple1hiddenlayer_results_rev02.docx")
print("Results saved to 03a_simple1hiddenlayer_results_rev02.docx") 