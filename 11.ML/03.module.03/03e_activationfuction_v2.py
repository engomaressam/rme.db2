# <p style="text-align:center">
#     <a href="https://skills.network" target="_blank">
#     <img src="https://cf-courses-data.s3.us.cloud-object-storage.appdomain.cloud/assets/logos/SN_web_lightmode.png" width="200" alt="Skills Network Logo">
#     </a>
# </p>

# <h1>Activation Functions</h1> 

# <h2>Objective</h2><ul><li> How to apply different Activation functions in Neural Network.</li></ul> 

# <h2>Table of Contents</h2>
# <p>In this lab, you will cover logistic regression by using PyTorch.</p>
#
# <ul>
#     <li><a href="#Log">Logistic Function</a></li>
#     <li><a href="#Tanh">Tanh</a></li>
#     <li><a href="#Relu">Relu</a></li>
#     <li><a href="#Compare">Compare Activation Functions</a></li>
# </ul>
# <p>Estimated Time Needed: <strong>15 min</strong></p>
#
# <hr>

# We'll need the following libraries
# Import the libraries we need for this lab
import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

output_buffer = io.StringIO()
plots = []

_original_show = plt.show
def save_and_collect(*args, **kwargs):
    fname = f"{os.path.splitext(os.path.basename(__file__))[0]}_plot_{len(plots)+1}.png"
    plt.savefig(fname)
    plots.append(fname)
    plt.close()
plt.show = save_and_collect

doc = Document()
doc.add_heading('Script Output', 0)

with redirect_stdout(output_buffer):
    # --- BEGIN MAIN SCRIPT ---
    import torch.nn as nn
    import torch
    import numpy as np
    torch.manual_seed(2)

    # <h2 id="Log">Logistic Function</h2>
    # Create a tensor ranging from -10 to 10: 
    # Create a tensor
    z = torch.arange(-10, 10, 0.1,).view(-1, 1)
    # When you use sequential, you can create a sigmoid object: 
    # Create a sigmoid object
    sig = nn.Sigmoid()
    # Apply the element-wise function Sigmoid with the object:
    # Make a prediction of sigmoid function
    yhat = sig(z)
    # Plot the results: 
    # Plot the result
    plt.plot(z.detach().numpy(),yhat.detach().numpy())
    plt.xlabel('z')
    plt.ylabel('yhat')
    # For custom modules, call the sigmoid from the torch (<code>nn.functional</code> for the old version), which applies the element-wise sigmoid from the function module and plots the results:
    # Use the build in function to predict the result
    yhat = torch.sigmoid(z)
    plt.plot(z.numpy(), yhat.numpy())

    # <h2 id="Tanh">Tanh</h2>
    # When you use sequential, you can create a tanh object:
    # Create a tanh object
    TANH = nn.Tanh()
    # Call the object and plot it:
    # Make the prediction using tanh object
    yhat = TANH(z)
    plt.plot(z.numpy(), yhat.numpy())
    # For custom modules, call the Tanh object from the torch (nn.functional for the old version), which applies the element-wise sigmoid from the function module and plots the results:
    # Make the prediction using the build-in tanh object
    yhat = torch.tanh(z)
    plt.plot(z.numpy(), yhat.numpy())

    # <h2 id="Relu">Relu</h2>
    # When you use sequential, you can create a Relu object: 
    # Create a relu object and make the prediction
    RELU = nn.ReLU()
    yhat = RELU(z)
    plt.plot(z.numpy(), yhat.numpy())
    # For custom modules, call the relu object from the nn.functional, which applies the element-wise sigmoid from the function module and plots the results:
    # Use the build-in function to make the prediction
    yhat = torch.relu(z)
    plt.plot(z.numpy(), yhat.numpy())

    # <a id="ref3"></a>
    # <h2> Compare Activation Functions </h2>
    # Plot the results to compare the activation functions
    x = torch.arange(-2, 2, 0.1).view(-1, 1)
    plt.plot(x.numpy(), torch.relu(x).numpy(), label='relu')
    plt.plot(x.numpy(), torch.sigmoid(x).numpy(), label='sigmoid')
    plt.plot(x.numpy(), torch.tanh(x).numpy(), label='tanh')
    plt.legend()

    # <a id="ref4"></a>
    # <h2> Practice </h2>
    # Compare the activation functions with a tensor in the range <i>(-1, 1)</i>
    # Practice: Compare the activation functions again using a tensor in the range (-1, 1)
    # Type your code here
    # Double-click <b>here</b> for the solution.
    # <!-- 
    # x = torch.arange(-1, 1, 0.1).view(-1, 1)
    # plt.plot(x.numpy(), torch.relu(x).numpy(), label = 'relu')
    # plt.plot(x.numpy(), torch.sigmoid(x).numpy(), label = 'sigmoid')
    # plt.plot(x.numpy(), torch.tanh(x).numpy(), label = 'tanh')
    # plt.legend()
    # -->
    # --- END MAIN SCRIPT ---

plt.show = _original_show
output_text = output_buffer.getvalue()
doc.add_paragraph(output_text)
for plot_file in plots:
    doc.add_picture(plot_file, width=Inches(5))
    os.remove(plot_file)
docx_name = f"{os.path.splitext(os.path.basename(__file__))[0]}.docx"
doc.save(docx_name)

# <h2>About the Authors:</h2> 
# <a href="https://www.linkedin.com/in/joseph-s-50398b136/">Joseph Santarcangelo</a> has a PhD in Electrical Engineering, his research focused on using machine learning, signal processing, and computer vision to determine how videos impact human cognition. Joseph has been working for IBM since he completed his PhD. 
# Other contributors: <a href="https://www.linkedin.com/in/michelleccarey/">Michelle Carey</a>, <a href="https://www.linkedin.com/in/jiahui-mavis-zhou-a4537814a">Mavis Zhou</a>
# <!--
# ## Change Log
#
# |  Date (YYYY-MM-DD) |  Version | Changed By  |  Change Description |
# |---|---|---|---|
# | 2020-09-23  | 2.0  | Shubham  |  Migrated Lab to Markdown and added to course repo in GitLab |
# -->
# <hr>
#
# ## <h3 align="center"> &#169; IBM Corporation. All rights reserved. <h3/> 