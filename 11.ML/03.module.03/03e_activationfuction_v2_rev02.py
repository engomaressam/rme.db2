import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import torch
import torch.nn as nn
import numpy as np

def main():
    torch.manual_seed(2025)
    z = torch.arange(-10, 10, 0.2)
    # Sigmoid
    yhat = torch.sigmoid(z)
    plt.plot(z.numpy(), yhat.numpy())
    plt.title('Sigmoid')
    plt.show()
    # Tanh
    yhat = torch.tanh(z)
    plt.plot(z.numpy(), yhat.numpy())
    plt.title('Tanh')
    plt.show()
    # ReLU
    yhat = torch.relu(z)
    plt.plot(z.numpy(), yhat.numpy())
    plt.title('ReLU')
    plt.show()
    # LeakyReLU
    leaky_relu = nn.LeakyReLU(0.1)
    yhat = leaky_relu(z)
    plt.plot(z.numpy(), yhat.numpy())
    plt.title('LeakyReLU')
    plt.show()
output_buffer = io.StringIO()
plots = []
_original_show = plt.show
def save_and_collect(*args, **kwargs):
    fname = f"{os.path.splitext(os.path.basename(__file__))[0]}_plot_{len(plots)+1}.png"
    plt.savefig(fname)
    plots.append(fname)
    plt.close()
plt.show = save_and_collect
doc = Document()
doc.add_heading('Script Output', 0)
with redirect_stdout(output_buffer):
    main()
plt.show = _original_show
output_text = output_buffer.getvalue()
doc.add_paragraph(output_text)
for plot_file in plots:
    doc.add_picture(plot_file, width=Inches(5))
    os.remove(plot_file)
docx_name = f"{os.path.splitext(os.path.basename(__file__))[0]}.docx"
doc.save(docx_name) 