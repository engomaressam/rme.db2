# <h1>Simple One Hidden Layer Neural Network</h1>
# <h2>Objective</h2><ul><li> How to create simple Neural Network in pytorch.</li></ul> 
# <h2>Table of Contents</h2>
# <p>In this lab, you will use a single-layer neural network to classify non linearly seprable data in 1-Ddatabase.</p>
# <ul>
#     <li><a href="#Model">Neural Network Module and Training Function</a></li>
#     <li><a href="#Makeup_Data">Make Some Data</a></li>
#     <li><a href="#Train">Define the Neural Network, Criterion Function, Optimizer, and Train the Model</a></li>
# </ul>
# <p>Estimated Time Needed: <strong>25 min</strong></p>
# <hr>
# <h2>Preparation</h2>
# We'll need the following libraries

# Import the libraries we need for this lab
import torch 
import torch.nn as nn
from torch import sigmoid
import matplotlib.pylab as plt
import numpy as np
import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
torch.manual_seed(0)

# Buffer for print output
output_buffer = io.StringIO()
plots = []

# Patch plt.show to save figures instead of displaying
_original_show = plt.show
def save_and_collect(*args, **kwargs):
    fname = f"{os.path.splitext(os.path.basename(__file__))[0]}_plot_{len(plots)+1}.png"
    plt.savefig(fname)
    plots.append(fname)
    plt.close()
plt.show = save_and_collect

doc = Document()
doc.add_heading('Script Output', 0)

with redirect_stdout(output_buffer):
    # Used for plotting the model
    # The function for plotting the model
    def PlotStuff(X, Y, model, epoch, leg=True):
        plt.plot(X.numpy(), model(X).detach().numpy(), label=('epoch ' + str(epoch)))
        plt.plot(X.numpy(), Y.numpy(), 'r')
        plt.xlabel('x')
        if leg == True:
            plt.legend()
        else:
            pass
    # <h2 id="Model">Neural Network Module and Training Function</h2> 
    # Define the activations and the output of the first linear layer as an attribute. Note that this is not good practice. 
    # Define the class Net
    class Net(nn.Module):
        # Constructor
        def __init__(self, D_in, H, D_out):
            super(Net, self).__init__()
            # hidden layer 
            self.linear1 = nn.Linear(D_in, H)
            self.linear2 = nn.Linear(H, D_out)
            # Define the first linear layer as an attribute, this is not good practice
            self.a1 = None
            self.l1 = None
            self.l2=None
        # Prediction
        def forward(self, x):
            self.l1 = self.linear1(x)
            self.a1 = sigmoid(self.l1)
            self.l2=self.linear2(self.a1)
            yhat = sigmoid(self.linear2(self.a1))
            return yhat
    # Define the training function:
    # Define the training function
    def train(Y, X, model, optimizer, criterion, epochs=1000):
        cost = []
        total=0
        for epoch in range(epochs):
            total=0
            for y, x in zip(Y, X):
                yhat = model(x)
                loss = criterion(yhat, y)
                loss.backward()
                optimizer.step()
                optimizer.zero_grad()
                #cumulative loss 
                total+=loss.item() 
            cost.append(total)
            if epoch % 300 == 0:    
                PlotStuff(X, Y, model, epoch, leg=True)
                plt.show()
                model(X)
                plt.scatter(model.a1.detach().numpy()[:, 0], model.a1.detach().numpy()[:, 1], c=Y.numpy().reshape(-1))
                plt.title('activations')
                plt.show()
        return cost
    # <h2 id="Makeup_Data">Make Some Data</h2>
    # Make some data
    X = torch.arange(-20, 20, 1).view(-1, 1).type(torch.FloatTensor)
    Y = torch.zeros(X.shape[0])
    Y[(X[:, 0] > -4) & (X[:, 0] < 4)] = 1.0
    # <h2 id="Train">Define the Neural Network, Criterion Function, Optimizer and Train the Model</h2>
    # Create the Cross-Entropy loss function: 
    def criterion_cross(outputs, labels):
        out = -1 * torch.mean(labels * torch.log(outputs) + (1 - labels) * torch.log(1 - outputs))
        return out
    # Define the Neural Network, Optimizer, and Train the Model:
    # Train the model
    # size of input 
    D_in = 1
    # size of hidden layer 
    H = 2
    # number of outputs 
    D_out = 1
    # learning rate 
    learning_rate = 0.1
    # create the model 
    model = Net(D_in, H, D_out)
    #optimizer 
    optimizer = torch.optim.SGD(model.parameters(), lr=learning_rate)
    #train the model usein
    cost_cross = train(Y, X, model, optimizer, criterion_cross, epochs=1000)
    #plot the loss
    plt.plot(cost_cross)
    plt.xlabel('epoch')
    plt.title('cross entropy loss')
    plt.show()
    # By examining the output of the  activation, you see by the 600th epoch that the data has been mapped to a linearly separable space.
    # we can make a prediction for a arbitrary one tensors 
    x=torch.tensor([0.0])
    yhat=model(x)
    print('Prediction for x=0.0:', yhat)
    # we can make a prediction for some arbitrary one tensors  
    X_=torch.tensor([[0.0],[2.0],[3.0]])
    Yhat=model(X_)
    print('Predictions for X_:', Yhat)
    # we  can threshold the predication
    Yhat=Yhat>0.5
    print('Thresholded predictions:', Yhat)
# Restore plt.show
plt.show = _original_show
# Write text output to docx
output_text = output_buffer.getvalue()
doc.add_paragraph(output_text)
# Add plots to docx
for plot_file in plots:
    doc.add_picture(plot_file, width=Inches(5))
    os.remove(plot_file)
# Save docx
docx_name = f"{os.path.splitext(os.path.basename(__file__))[0]}.docx"
doc.save(docx_name)