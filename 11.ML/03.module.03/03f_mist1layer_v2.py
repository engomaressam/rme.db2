# <p style="text-align:center">
#     <a href="https://skills.network" target="_blank">
#     <img src="https://cf-courses-data.s3.us.cloud-object-storage.appdomain.cloud/assets/logos/SN_web_lightmode.png" width="200" alt="Skills Network Logo">
#     </a>
# </p>

# <h1>Test Sigmoid, Tanh, and Relu Activations Functions on the MNIST Dataset</h1>

# <h2>Objective</h2><ul><li> How to apply different activation functions on the MNIST dataset.</li></ul> 

# <h2>Table of Contents</h2>
# <p>In this lab, you will test sigmoid, tanh, and relu activation functions on the MNIST dataset.</p>
#
# <ul>
#     <li><a href="#Model">Neural Network Module and Training Function</a></li>
#     <li><a href="#Makeup_Data">Make Some Data</a></li>
#     <li><a href="#Train">Define Several Neural Network, Criterion Function, and Optimizer</a></li>
#     <li><a href="#Test">Test Sigmoid, Tanh, and Relu</a></li>
#     <li><a href="#Result">Analyze Results</a></li>
# </ul>
# <p></p>
# Estimated Time Needed: <strong>25 min</strong>
# </div>
#
# <hr>

# <h2>Preparation</h2>
# We'll need the following libraries
# Uncomment the following line to install the torchvision library
# !mamba install -y torchvision
# Import the libraries we need for this lab
import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

output_buffer = io.StringIO()
plots = []

_original_show = plt.show
def save_and_collect(*args, **kwargs):
    fname = f"{os.path.splitext(os.path.basename(__file__))[0]}_plot_{len(plots)+1}.png"
    plt.savefig(fname)
    plots.append(fname)
    plt.close()
plt.show = save_and_collect

doc = Document()
doc.add_heading('Script Output', 0)

with redirect_stdout(output_buffer):
    # --- BEGIN MAIN SCRIPT ---
    import torch
    import torch.nn as nn
    import torchvision.transforms as transforms
    import torchvision.datasets as dsets
    import numpy as np
    torch.manual_seed(0)

    # <h2 id="Model">Neural Network Module and Training Function</h2> 
    # Define the neural network module or class using the sigmoid activation function: 
    # Build the model with sigmoid function
    class Net(nn.Module):
        # Constructor
        def __init__(self, D_in, H, D_out):
            super(Net, self).__init__()
            self.linear1 = nn.Linear(D_in, H)
            self.linear2 = nn.Linear(H, D_out)
        # Prediction
        def forward(self, x):
            x = torch.sigmoid(self.linear1(x))  
            x = self.linear2(x)
            return x

    # Define the neural network module or class using the Tanh activation function:
    # Build the model with Tanh function
    class NetTanh(nn.Module):
        # Constructor
        def __init__(self, D_in, H, D_out):
            super(NetTanh, self).__init__()
            self.linear1 = nn.Linear(D_in, H)
            self.linear2 = nn.Linear(H, D_out)
        # Prediction
        def forward(self, x):
            x = torch.tanh(self.linear1(x))
            x = self.linear2(x)
            return x

    # Define the neural network module or class using the Relu activation function:
    # Build the model with Relu function
    class NetRelu(nn.Module):
        # Constructor
        def __init__(self, D_in, H, D_out):
            super(NetRelu, self).__init__()
            self.linear1 = nn.Linear(D_in, H)
            self.linear2 = nn.Linear(H, D_out)
        # Prediction
        def forward(self, x):
            x = torch.relu(self.linear1(x))
            x = self.linear2(x)
            return x

    # Define a function to train the model. In this case, the function returns a Python dictionary to store the training loss for each iteration  and accuracy on the validation data.
    def train(model, criterion, train_loader, validation_loader, optimizer, epochs = 100):
        i = 0
        useful_stuff = {'training_loss':[], 'validation_accuracy':[]}  
        for epoch in range(epochs):
            for i, (x, y) in enumerate(train_loader):
                optimizer.zero_grad()
                z = model(x.view(-1, 28 * 28))
                loss = criterion(z, y)
                loss.backward()
                optimizer.step()
                useful_stuff['training_loss'].append(loss.item())
            correct = 0
            for x, y in validation_loader:
                z = model(x.view(-1, 28 * 28))
                _, label=torch.max(z, 1)
                correct += (label == y).sum().item()
            accuracy = 100 * (correct / len(validation_dataset))
            useful_stuff['validation_accuracy'].append(accuracy)
        return useful_stuff

    # <h2 id="Makeup_Data">Make Some Data</h2> 
    # Load the training dataset by setting the parameters <code>train</code> to <code>True</code> and convert it to a tensor by placing a transform object in the argument <code>transform</code>.
    # Create the training dataset
    train_dataset = dsets.MNIST(root='./data', train=True, download=True, transform=transforms.ToTensor())
    # Load the testing dataset by setting the parameter <code>train</code> to <code>False</code> and convert it to a tensor by placing a transform object in the argument <code>transform</code>.
    # Create the validation  dataset
    validation_dataset = dsets.MNIST(root='./data', train=False, download=True, transform=transforms.ToTensor())
    # Create the criterion function:  
    # Create the criterion function
    criterion = nn.CrossEntropyLoss()
    # Create the training-data loader and the validation-data loader object:
    # Create the training data loader and validation data loader object
    train_loader = torch.utils.data.DataLoader(dataset=train_dataset, batch_size=2000, shuffle=True)
    validation_loader = torch.utils.data.DataLoader(dataset=validation_dataset, batch_size=5000, shuffle=False)
    # <h2 id="Train">Define the Neural Network, Criterion Function, Optimizer, and Train the Model</h2> 
    # Create the criterion function: 
    # Create the criterion function
    criterion = nn.CrossEntropyLoss()
    # Create the model with 100 hidden neurons:  
    # Create the model object
    input_dim = 28 * 28
    hidden_dim = 100
    output_dim = 10
    model = Net(input_dim, hidden_dim, output_dim)
    # <h2 id="Test">Test Sigmoid, Tanh, and Relu</h2> 
    # Train the network by using the sigmoid activations function:
    # Train a model with sigmoid function
    learning_rate = 0.01
    optimizer = torch.optim.SGD(model.parameters(), lr=learning_rate)
    training_results = train(model, criterion, train_loader, validation_loader, optimizer, epochs=30)
    # Train the network by using the Tanh activations function:
    # Train a model with Tanh function
    model_Tanh = NetTanh(input_dim, hidden_dim, output_dim)
    optimizer = torch.optim.SGD(model_Tanh.parameters(), lr=learning_rate)
    training_results_tanch = train(model_Tanh, criterion, train_loader, validation_loader, optimizer, epochs=30)
    # Train the network by using the Relu activations function:
    # Train a model with Relu function
    modelRelu = NetRelu(input_dim, hidden_dim, output_dim)
    optimizer = torch.optim.SGD(modelRelu.parameters(), lr=learning_rate)
    training_results_relu = train(modelRelu, criterion, train_loader, validation_loader, optimizer, epochs=30)
    # <h2 id="Result">Analyze Results</h2> 
    # Compare the training loss for each activation: 
    # Compare the training loss
    plt.plot(training_results_tanch['training_loss'], label='tanh')
    plt.plot(training_results['training_loss'], label='sigmoid')
    plt.plot(training_results_relu['training_loss'], label='relu')
    plt.ylabel('loss')
    plt.title('training loss iterations')
    plt.legend()
    plt.show()
    # Compare the validation loss for each model:  
    # Compare the validation loss
    plt.plot(training_results_tanch['validation_accuracy'], label='tanh')
    plt.plot(training_results['validation_accuracy'], label='sigmoid')
    plt.plot(training_results_relu['validation_accuracy'], label='relu') 
    plt.ylabel('validation accuracy')
    plt.xlabel('epochs ')
    plt.legend()
    plt.show()
    # ## Which activation function performed best ?
    # <h2>About the Authors:</h2> 
    # <a href="https://www.linkedin.com/in/joseph-s-50398b136/">Joseph Santarcangelo</a> has a PhD in Electrical Engineering, his research focused on using machine learning, signal processing, and computer vision to determine how videos impact human cognition. Joseph has been working for IBM since he completed his PhD. 
    # Other contributors: <a href="https://www.linkedin.com/in/michelleccarey/">Michelle Carey</a>, <a href="https://www.linkedin.com/in/jiahui-mavis-zhou-a4537814a">Mavis Zhou</a>
    # <!--
    # ## Change Log
    #
    # |  Date (YYYY-MM-DD) |  Version | Changed By  |  Change Description |
    # |---|---|---|---|
    # | 2020-09-23  | 2.0  | Shubham  |  Migrated Lab to Markdown and added to course repo in GitLab |
    # -->
    # <hr>
    #
    # ## <h3 align="center"> &#169; IBM Corporation. All rights reserved. <h3/> 
    # --- END MAIN SCRIPT ---

plt.show = _original_show
output_text = output_buffer.getvalue()
doc.add_paragraph(output_text)
for plot_file in plots:
    doc.add_picture(plot_file, width=Inches(5))
    os.remove(plot_file)
docx_name = f"{os.path.splitext(os.path.basename(__file__))[0]}.docx"
doc.save(docx_name) 