import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import torch
import torch.nn as nn
import torchvision.transforms as transforms
import torchvision.datasets as dsets
import numpy as np

def main():
    torch.manual_seed(2026)
    input_dim = 28 * 28
    hidden_dim = 64
    output_dim = 10
    learning_rate = 0.002
    batch_size = 128
    epochs = 2
    train_dataset = dsets.MNIST(root='./data', train=True, transform=transforms.ToTensor(), download=True)
    test_dataset = dsets.MNIST(root='./data', train=False, transform=transforms.ToTensor(), download=True)
    train_loader = torch.utils.data.DataLoader(dataset=train_dataset, batch_size=batch_size, shuffle=True)
    test_loader = torch.utils.data.DataLoader(dataset=test_dataset, batch_size=batch_size, shuffle=False)
    class Net(nn.Module):
        def __init__(self, input_dim, hidden_dim, output_dim, activation):
            super(Net, self).__init__()
            self.fc1 = nn.Linear(input_dim, hidden_dim)
            self.fc2 = nn.Linear(hidden_dim, output_dim)
            self.activation = activation
        def forward(self, x):
            x = x.view(-1, input_dim)
            x = self.activation(self.fc1(x))
            x = self.fc2(x)
            return x
    activations = [torch.relu, torch.sigmoid, torch.tanh]
    for act in activations:
        model = Net(input_dim, hidden_dim, output_dim, act)
        criterion = nn.CrossEntropyLoss()
        optimizer = torch.optim.Adam(model.parameters(), lr=learning_rate)
        for epoch in range(epochs):
            total_loss = 0
            for images, labels in train_loader:
                outputs = model(images)
                loss = criterion(outputs, labels)
                optimizer.zero_grad()
                loss.backward()
                optimizer.step()
                total_loss += loss.item()
            print(f"Activation: {act.__name__}, Epoch [{epoch+1}/{epochs}], Loss: {total_loss:.4f}")
        # Evaluate
        model.eval()
        correct = 0
        total = 0
        with torch.no_grad():
            for images, labels in test_loader:
                outputs = model(images)
                _, predicted = torch.max(outputs.data, 1)
                total += labels.size(0)
                correct += (predicted == labels).sum().item()
        print(f'Activation: {act.__name__}, Accuracy on 10000 test images: {100 * correct / total:.2f}%')
output_buffer = io.StringIO()
plots = []
_original_show = plt.show
def save_and_collect(*args, **kwargs):
    fname = f"{os.path.splitext(os.path.basename(__file__))[0]}_plot_{len(plots)+1}.png"
    plt.savefig(fname)
    plots.append(fname)
    plt.close()
plt.show = save_and_collect
doc = Document()
doc.add_heading('Script Output', 0)
with redirect_stdout(output_buffer):
    main()
plt.show = _original_show
output_text = output_buffer.getvalue()
doc.add_paragraph(output_text)
for plot_file in plots:
    doc.add_picture(plot_file, width=Inches(5))
    os.remove(plot_file)
docx_name = f"{os.path.splitext(os.path.basename(__file__))[0]}.docx"
doc.save(docx_name) 