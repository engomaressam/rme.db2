import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import torch
import torch.nn as nn
import numpy as np
from torch.utils.data import Dataset, DataLoader

def main():
    torch.manual_seed(99)
    class Net(nn.Module):
        def __init__(self, D_in, H, D_out):
            super(Net, self).__init__()
            self.linear1 = nn.Linear(D_in, H)
            self.linear2 = nn.Linear(H, D_out)
        def forward(self, x):
            x = torch.sigmoid(self.linear1(x))
            x = torch.sigmoid(self.linear2(x))
            return x
    def train(data_set, model, criterion, train_loader, optimizer, epochs=300):
        COST = []
        ACC = []
        for epoch in range(epochs):
            total = 0
            for x, y in train_loader:
                optimizer.zero_grad()
                yhat = model(x)
                loss = criterion(yhat, y)
                optimizer.zero_grad()
                loss.backward()
                optimizer.step()
                total += loss.item()
            ACC.append(np.mean(data_set.y.view(-1).numpy() == (model(data_set.x)[:, 0] > 0.5).numpy()))
            COST.append(total)
        fig, ax1 = plt.subplots()
        color = 'tab:red'
        ax1.plot(COST, color=color)
        ax1.set_xlabel('epoch', color=color)
        ax1.set_ylabel('total loss', color=color)
        ax1.tick_params(axis='y', color=color)
        ax2 = ax1.twinx()
        color = 'tab:blue'
        ax2.set_ylabel('accuracy', color=color)
        ax2.plot(ACC, color=color)
        ax2.tick_params(axis='y', color=color)
        fig.tight_layout()
        plt.show()
        return COST
    class XOR_Data(Dataset):
        def __init__(self, N_s=200):
            self.x = torch.zeros((N_s, 2))
            self.y = torch.zeros((N_s, 1))
            for i in range(N_s // 4):
                self.x[i, :] = torch.Tensor([0.0, 0.0])
                self.y[i, 0] = torch.Tensor([0.0])
                self.x[i + N_s // 4, :] = torch.Tensor([0.0, 1.0])
                self.y[i + N_s // 4, 0] = torch.Tensor([1.0])
                self.x[i + N_s // 2, :] = torch.Tensor([1.0, 0.0])
                self.y[i + N_s // 2, 0] = torch.Tensor([1.0])
                self.x[i + 3 * N_s // 4, :] = torch.Tensor([1.0, 1.0])
                self.y[i + 3 * N_s // 4, 0] = torch.Tensor([0.0])
                self.x = self.x + 0.02 * torch.randn((N_s, 2))
            self.len = N_s
        def __getitem__(self, index):
            return self.x[index], self.y[index]
        def __len__(self):
            return self.len
        def plot_stuff(self):
            plt.plot(self.x[self.y[:, 0] == 0, 0].numpy(), self.x[self.y[:, 0] == 0, 1].numpy(), 'o', label="y=0")
            plt.plot(self.x[self.y[:, 0] == 1, 0].numpy(), self.x[self.y[:, 0] == 1, 1].numpy(), 'ro', label="y=1")
            plt.legend()
    data_set = XOR_Data()
    data_set.plot_stuff()
    # One neuron
    model = Net(2, 1, 1)
    learning_rate = 0.005
    criterion = nn.BCELoss()
    optimizer = torch.optim.Adam(model.parameters(), lr=learning_rate)
    train_loader = DataLoader(dataset=data_set, batch_size=4)
    LOSS12 = train(data_set, model, criterion, train_loader, optimizer, epochs=300)
    plt.title('One Neuron')
    plt.show()
    # Two neurons
    model = Net(2, 2, 1)
    learning_rate = 0.02
    criterion = nn.BCELoss()
    optimizer = torch.optim.Adam(model.parameters(), lr=learning_rate)
    train_loader = DataLoader(dataset=data_set, batch_size=4)
    LOSS12 = train(data_set, model, criterion, train_loader, optimizer, epochs=300)
    plt.title('Two Neurons')
    plt.show()
    # Three neurons
    model = Net(2, 3, 1)
    learning_rate = 0.02
    criterion = nn.BCELoss()
    optimizer = torch.optim.Adam(model.parameters(), lr=learning_rate)
    train_loader = DataLoader(dataset=data_set, batch_size=4)
    LOSS12 = train(data_set, model, criterion, train_loader, optimizer, epochs=300)
    plt.title('Three Neurons')
    plt.show()
output_buffer = io.StringIO()
plots = []
_original_show = plt.show
def save_and_collect(*args, **kwargs):
    fname = f"{os.path.splitext(os.path.basename(__file__))[0]}_plot_{len(plots)+1}.png"
    plt.savefig(fname)
    plots.append(fname)
    plt.close()
plt.show = save_and_collect
doc = Document()
doc.add_heading('Script Output', 0)
with redirect_stdout(output_buffer):
    main()
plt.show = _original_show
output_text = output_buffer.getvalue()
doc.add_paragraph(output_text)
for plot_file in plots:
    doc.add_picture(plot_file, width=Inches(5))
    os.remove(plot_file)
docx_name = f"{os.path.splitext(os.path.basename(__file__))[0]}.docx"
doc.save(docx_name) 