# <h1>Softmax Classifier 1D - rev.02</h1>
# Objective: Build a Softmax classifier using the Sequential module in PyTorch with different parameters than the original notebook.
# This script uses the GPU if available and saves all results and figures to a Word file.

import torch
import torch.nn as nn
import matplotlib.pyplot as plt
import numpy as np
from torch.utils.data import Dataset, DataLoader
import os
from docx import Document
from docx.shared import Inches

device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
torch.manual_seed(42)

# Helper function to plot labeled data points

def plot_data(data_set, model=None, n=1, color=False, save_path=None):
    X = data_set[:][0]
    Y = data_set[:][1]
    plt.plot(X[Y == 0, 0].cpu().numpy(), Y[Y == 0].cpu().numpy(), 'bo', label='y = 0')
    plt.plot(X[Y == 1, 0].cpu().numpy(), 0 * Y[Y == 1].cpu().numpy(), 'ro', label='y = 1')
    plt.plot(X[Y == 2, 0].cpu().numpy(), 0 * Y[Y == 2].cpu().numpy(), 'go', label='y = 2')
    plt.ylim((-0.1, 3))
    plt.legend()
    if model is not None:
        w = list(model.parameters())[0][0].detach()
        b = list(model.parameters())[1][0].detach()
        y_label = ['yhat=0', 'yhat=1', 'yhat=2']
        y_color = ['b', 'r', 'g']
        Ylines = []
        for w, b, y_l, y_c in zip(model.state_dict()['0.weight'], model.state_dict()['0.bias'], y_label, y_color):
            Ylines.append((w * X + b).cpu().numpy())
            plt.plot(X.cpu().numpy(), (w * X + b).cpu().numpy(), y_c, label=y_l)
    plt.legend()
    if save_path:
        plt.savefig(save_path)
    plt.close()

# Data class with different parameters
class Data(Dataset):
    def __init__(self):
        self.x = torch.arange(-3, 3, 0.2).view(-1, 1)
        self.y = torch.zeros(self.x.shape[0])
        self.y[(self.x > -1.5)[:, 0] * (self.x < 1.5)[:, 0]] = 1
        self.y[(self.x >= 1.5)[:, 0]] = 2
        self.y = self.y.type(torch.LongTensor)
        self.len = self.x.shape[0]
    def __getitem__(self, index):
        return self.x[index].to(device), self.y[index].to(device)
    def __len__(self):
        return self.len

data_set = Data()

# Plot and save the dataset
os.makedirs('figs_softmax_rev02', exist_ok=True)
plot_path = 'figs_softmax_rev02/dataset.png'
plot_data(data_set, save_path=plot_path)

doc = Document()
doc.add_heading('Softmax Classifier 1D - rev.02', 0)
doc.add_paragraph(f"Device used: {device}")
doc.add_paragraph("Dataset plot:")
doc.add_picture(plot_path, width=Inches(4))

# Build Softmax Classifier
model = nn.Sequential(nn.Linear(1, 3)).to(device)

# Criterion, optimizer, dataloader with different parameters
criterion = nn.CrossEntropyLoss()
optimizer = torch.optim.Adam(model.parameters(), lr=0.05)
trainloader = DataLoader(dataset=data_set, batch_size=8)

# Train the model with different epochs
LOSS = []
def train_model(epochs):
    for epoch in range(epochs):
        if epoch % 40 == 0:
            plot_path = f'figs_softmax_rev02/model_epoch_{epoch}.png'
            plot_data(data_set, model, save_path=plot_path)
            doc.add_paragraph(f"Model output at epoch {epoch}")
            doc.add_picture(plot_path, width=Inches(4))
        for x, y in trainloader:
            optimizer.zero_grad()
            yhat = model(x)
            loss = criterion(yhat, y)
            LOSS.append(loss.item())
            loss.backward()
            optimizer.step()
train_model(200)

# Plot and save the loss curve
plt.figure()
plt.plot(LOSS)
plt.xlabel('iteration')
plt.title('Loss curve')
loss_fig_path = 'figs_softmax_rev02/loss_curve.png'
plt.savefig(loss_fig_path)
plt.close()
doc.add_paragraph("Loss curve:")
doc.add_picture(loss_fig_path, width=Inches(4))

# Make the prediction
z = model(data_set.x.to(device))
_, yhat = z.max(1)
doc.add_paragraph(f"The prediction: {yhat.cpu().numpy()}")

# Print the accuracy
correct = (data_set.y.to(device) == yhat).sum().item()
accuracy = correct / len(data_set)
doc.add_paragraph(f"The accuracy: {accuracy}")

# Softmax probabilities
Softmax_fn = nn.Softmax(dim=-1)
Probability = Softmax_fn(z)
for i in range(3):
    doc.add_paragraph(f"probability of class {i} for first sample: {Probability[0,i].item()}")

doc.save("02a_softmax_one_dimension_results_rev02.docx")
print("Results saved to 02a_softmax_one_dimension_results_rev02.docx") 