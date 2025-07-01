"""
Convert Jupyter notebooks to Python scripts with GPU monitoring and save outputs to Word
"""
import os
import json
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Pt

def add_gpu_monitoring(cells):
    """Add GPU monitoring code to the notebook cells"""
    gpu_check = {
        "cell_type": "code",
        "execution_count": None,
        "metadata": {},
        "outputs": [],
        "source": [
            "# GPU Monitoring Setup\n",
            "import torch\n",
            "import subprocess\n",
            "from datetime import datetime\n",
            "print(\"="*80)\n",
            "print(\"GPU STATUS\")\n",
            "print(\"="*80)\n",
            "print(f\"PyTorch Version: {torch.__version__}\")\n",
            "print(f\"CUDA Available: {torch.cuda.is_available()}\")\n",
            "if torch.cuda.is_available():\n",
            "    print(f\"CUDA Version: {torch.version.cuda}\")\n",
            "    print(f\"GPU: {torch.cuda.get_device_name(0)}\")\n",
            "    print(f\"Memory Allocated: {torch.cuda.memory_allocated()/1024**2:.2f} MB\")\n",
            "    print(f\"Memory Cached: {torch.cuda.memory_reserved()/1024**2:.2f} MB\")\n",
            "print(\"="*80)\n"
        ]
    }
    
    # Add GPU check at the beginning
    cells.insert(1, gpu_check)
    
    return cells

def convert_notebook(notebook_path, output_dir):
    """Convert notebook to Python script with GPU monitoring"""
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Load notebook
    with open(notebook_path, 'r', encoding='utf-8') as f:
        notebook = json.load(f)
    
    # Add GPU monitoring
    notebook['cells'] = add_gpu_monitoring(notebook['cells'])
    
    # Save modified notebook
    temp_nb = os.path.join(output_dir, 'temp.ipynb')
    with open(temp_nb, 'w', encoding='utf-8') as f:
        json.dump(notebook, f)
    
    # Convert to Python script
    script_name = os.path.splitext(os.path.basename(notebook_path))[0] + '.py'
    script_path = os.path.join(output_dir, script_name)
    subprocess.run(['jupyter', 'nbconvert', '--to', 'script', temp_nb, '--output', script_path])
    
    # Remove temporary notebook
    os.remove(temp_nb)
    
    return script_path

def run_script(script_path, output_dir):
    """Run Python script and capture output to Word document"""
    # Create output filename
    base_name = os.path.splitext(os.path.basename(script_path))[0]
    docx_path = os.path.join(output_dir, f"{base_name}_output.docx")
    
    # Run script and capture output
    result = subprocess.run(['python', script_path], 
                          capture_output=True, 
                          text=True)
    
    # Create Word document
    doc = Document()
    doc.add_heading(f'Output of {base_name}', 0)
    
    # Add timestamp
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Add output
    doc.add_heading('Standard Output:', level=1)
    doc.add_paragraph(result.stdout)
    
    if result.stderr:
        doc.add_heading('Errors:', level=1)
        doc.add_paragraph(result.stderr)
    
    # Save document
    doc.save(docx_path)
    print(f"Output saved to: {docx_path}")

def main():
    # Set up directories
    base_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(base_dir, 'notebook_outputs')
    os.makedirs(output_dir, exist_ok=True)
    
    # Process each notebook
    notebooks = [
        '01.cross_entropy_logistic_regression_v2.ipynb',
        '02.softmax_in_one_dimension_v2.ipynb',
        '03.lab_predicting _MNIST_using_Softmax_v2.ipynb'
    ]
    
    for nb in notebooks:
        nb_path = os.path.join(base_dir, nb)
        if os.path.exists(nb_path):
            try:
                print(f"\nProcessing {nb}...")
                script_path = convert_notebook(nb_path, output_dir)
                run_script(script_path, output_dir)
            except Exception as e:
                print(f"Error processing {nb}: {str(e)}")
        else:
            print(f"Notebook not found: {nb}")

if __name__ == "__main__":
    main()
