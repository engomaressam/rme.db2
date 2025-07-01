# <h1>Softmax Classifier MNIST - rev.02</h1>
# Objective: Classify handwritten digits from the MNIST database using a Softmax classifier with different parameters than the original notebook.
# This script uses the GPU if available and saves all results and figures to a Word file.

import torch
import torch.nn as nn
import torchvision.transforms as transforms
import torchvision.datasets as dsets
import matplotlib.pyplot as plt
import numpy as np
import os
from docx import Document
from docx.shared import Inches

device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
torch.manual_seed(123)

# Helper function to plot parameters

def PlotParameters(model, save_path=None):
    W = model.state_dict()['linear.weight'].data.cpu()
    w_min = W.min().item()
    w_max = W.max().item()
    fig, axes = plt.subplots(2, 5)
    fig.subplots_adjust(hspace=0.01, wspace=0.1)
    for i, ax in enumerate(axes.flat):
        if i < 10:
            ax.set_xlabel(f"class: {i}")
            ax.imshow(W[i, :].view(28, 28), vmin=w_min, vmax=w_max, cmap='seismic')
            ax.set_xticks([])
            ax.set_yticks([])
    if save_path:
        plt.savefig(save_path)
    plt.close()

# Helper function to show data

def show_data(data_sample, save_path=None):
    plt.imshow(data_sample[0].cpu().numpy().reshape(28, 28), cmap='gray')
    plt.title('y = ' + str(data_sample[1]))
    if save_path:
        plt.savefig(save_path)
    plt.close()

# Load datasets
train_dataset = dsets.MNIST(root='./data', train=True, download=True, transform=transforms.ToTensor())
validation_dataset = dsets.MNIST(root='./data', download=True, transform=transforms.ToTensor())

# Model definition
class SoftMax(nn.Module):
    def __init__(self, input_size, output_size):
        super(SoftMax, self).__init__()
        self.linear = nn.Linear(input_size, output_size)
    def forward(self, x):
        z = self.linear(x)
        return z

input_dim = 28 * 28
output_dim = 10
model = SoftMax(input_dim, output_dim).to(device)

# Different parameters
learning_rate = 0.02
optimizer = torch.optim.Adam(model.parameters(), lr=learning_rate)
criterion = nn.CrossEntropyLoss()
train_loader = torch.utils.data.DataLoader(dataset=train_dataset, batch_size=256)
validation_loader = torch.utils.data.DataLoader(dataset=validation_dataset, batch_size=10000)

n_epochs = 5
loss_list = []
accuracy_list = []
N_test = len(validation_dataset)

# Prepare docx
doc = Document()
doc.add_heading('Softmax Classifier MNIST - rev.02', 0)
doc.add_paragraph(f"Device used: {device}")

# Plot initial parameters
os.makedirs('figs_mnist_softmax_rev02', exist_ok=True)
param_fig_path = 'figs_mnist_softmax_rev02/params_initial.png'
PlotParameters(model, save_path=param_fig_path)
doc.add_paragraph("Initial model parameters:")
doc.add_picture(param_fig_path, width=Inches(4))

def train_model(n_epochs):
    for epoch in range(n_epochs):
        for x, y in train_loader:
            x, y = x.to(device), y.to(device)
            optimizer.zero_grad()
            z = model(x.view(-1, 28 * 28))
            loss = criterion(z, y)
            loss.backward()
            optimizer.step()
        correct = 0
        for x_test, y_test in validation_loader:
            x_test, y_test = x_test.to(device), y_test.to(device)
            z = model(x_test.view(-1, 28 * 28))
            _, yhat = torch.max(z.data, 1)
            correct += (yhat == y_test).sum().item()
        accuracy = correct / N_test
        loss_list.append(loss.data.cpu().item())
        accuracy_list.append(accuracy)
        # Save loss/accuracy plot every epoch
        if (epoch+1) % 1 == 0:
            plt.figure()
            plt.plot(loss_list, label='Loss')
            plt.plot(accuracy_list, label='Accuracy')
            plt.xlabel('epoch')
            plt.legend()
            plt.title('Loss and Accuracy')
            plot_path = f'figs_mnist_softmax_rev02/loss_acc_epoch_{epoch+1}.png'
            plt.savefig(plot_path)
            plt.close()
            doc.add_paragraph(f"Loss and accuracy at epoch {epoch+1}:")
            doc.add_picture(plot_path, width=Inches(4))
train_model(n_epochs)

# Plot final parameters
param_fig_path_final = 'figs_mnist_softmax_rev02/params_final.png'
PlotParameters(model, save_path=param_fig_path_final)
doc.add_paragraph("Final model parameters:")
doc.add_picture(param_fig_path_final, width=Inches(4))

doc.save("02b_mnist_softmax_results_rev02.docx")
print("Results saved to 02b_mnist_softmax_results_rev02.docx") 