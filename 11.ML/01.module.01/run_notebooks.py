"""
Run all notebooks and save outputs to Word documents
"""
import os
import subprocess
from datetime import datetime
from docx import Document

def run_notebook(notebook_name):
    """Run a notebook and save output to Word"""
    print(f"Processing {notebook_name}...")
    
    # Create output directory
    output_dir = "notebook_outputs"
    os.makedirs(output_dir, exist_ok=True)
    
    # Create output filename
    base_name = os.path.splitext(notebook_name)[0]
    output_file = os.path.join(output_dir, f"{base_name}_output.docx")
    
    # Add GPU monitoring to notebook
    with open(notebook_name, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Add GPU check at the beginning
    gpu_check = """
# GPU Monitoring Setup
import torch
import subprocess
print("="*80)
print("GPU STATUS")
print("="*80)
print(f"PyTorch Version: {torch.__version__}")
print(f"CUDA Available: {torch.cuda.is_available()}")
if torch.cuda.is_available():
    print(f"CUDA Version: {torch.version.cuda}")
    print(f"GPU: {torch.cuda.get_device_name(0)}")
    print(f"Memory Allocated: {torch.cuda.memory_allocated()/1024**2:.2f} MB")
    print(f"Memory Cached: {torch.cuda.memory_reserved()/1024**2:.2f} MB")
print("="*80)
"""
    
    # Create temp notebook with GPU monitoring
    temp_nb = "temp_notebook.ipynb"
    with open(temp_nb, 'w', encoding='utf-8') as f:
        f.write(content.replace('"cells": [', '"cells": [\n    {'
            f'"cell_type": "code",\n     "execution_count": null,\n     "metadata": {{}},\n     "outputs": [],\n     "source": ["{gpu_check.replace(chr(10), chr(92)+chr(110) + \"\\" + \"\" + chr(34) + \" \" + chr(34)}\"]\n    },'))
    
    try:
        # Run the notebook
        result = subprocess.run(['jupyter', 'nbconvert', '--execute', '--to', 'notebook', '--inplace', temp_nb],
                              capture_output=True, text=True)
        
        # Convert to HTML
        subprocess.run(['jupyter', 'nbconvert', '--to', 'html', temp_nb, '--output', os.path.join(output_dir, base_name)])
        
        # Convert HTML to Word
        doc = Document()
        doc.add_heading(f'Output of {notebook_name}', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add output
        with open(os.path.join(output_dir, f"{base_name}.html"), 'r', encoding='utf-8') as f:
            html_content = f.read()
        doc.add_paragraph(html_content)
        
        doc.save(output_file)
        print(f"Output saved to: {output_file}")
        
    except Exception as e:
        print(f"Error processing {notebook_name}: {str(e)}")
    finally:
        # Clean up
        if os.path.exists(temp_nb):
            os.remove(temp_nb)
        html_file = os.path.join(output_dir, f"{base_name}.html")
        if os.path.exists(html_file):
            os.remove(html_file)

def main():
    # List of notebooks to process
    notebooks = [
        '01.cross_entropy_logistic_regression_v2.ipynb',
        '02.softmax_in_one_dimension_v2.ipynb',
        '03.lab_predicting _MNIST_using_Softmax_v2.ipynb'
    ]
    
    # Process each notebook
    for nb in notebooks:
        if os.path.exists(nb):
            run_notebook(nb)
        else:
            print(f"Notebook not found: {nb}")

if __name__ == "__main__":
    main()
