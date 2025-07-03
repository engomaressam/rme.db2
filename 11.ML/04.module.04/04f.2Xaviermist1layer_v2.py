import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

script_name = os.path.splitext(os.path.basename(__file__))[0]
word_output = f"{script_name}.docx"
doc = Document()
doc.add_heading(f'Output for {script_name}', 0)

f = io.StringIO()
with redirect_stdout(f):
    # --- Begin original script ---
    import torch
    import torch.nn as nn
    import torchvision.transforms as transforms
    import torchvision.datasets as dsets
    import matplotlib.pylab as plt
    import numpy as np
    torch.manual_seed(0)

    class Net_Xavier(nn.Module):
        def __init__(self, Layers):
            super(Net_Xavier, self).__init__()
            self.hidden = nn.ModuleList()
            for input_size, output_size in zip(Layers, Layers[1:]):
                linear = nn.Linear(input_size, output_size)
                torch.nn.init.xavier_uniform_(linear.weight)
                self.hidden.append(linear)
        def forward(self, x):
            L = len(self.hidden)
            for (l, linear_transform) in zip(range(L), self.hidden):
                if l < L - 1:
                    x = torch.tanh(linear_transform(x))
                else:
                    x = linear_transform(x)
            return x

    class Net_Uniform(nn.Module):
        def __init__(self, Layers):
            super(Net_Uniform, self).__init__()
            self.hidden = nn.ModuleList()
            for input_size, output_size in zip(Layers, Layers[1:]):
                linear = nn.Linear(input_size, output_size)
                linear.weight.data.uniform_(0, 1)
                self.hidden.append(linear)
        def forward(self, x):
            L = len(self.hidden)
            for (l, linear_transform) in zip(range(L), self.hidden):
                if l < L - 1:
                    x = torch.tanh(linear_transform(x))
                else:
                    x = linear_transform(x)
            return x

    class Net(nn.Module):
        def __init__(self, Layers):
            super(Net, self).__init__()
            self.hidden = nn.ModuleList()
            for input_size, output_size in zip(Layers, Layers[1:]):
                linear = nn.Linear(input_size, output_size)
                self.hidden.append(linear)
        def forward(self, x):
            L = len(self.hidden)
            for (l, linear_transform) in zip(range(L), self.hidden):
                if l < L - 1:
                    x = torch.tanh(linear_transform(x))
                else:
                    x = linear_transform(x)
            return x

    def train(model, criterion, train_loader, validation_loader, optimizer, epochs = 100):
        loss_accuracy = {'training_loss':[], 'validation_accuracy':[]}
        for epoch in range(epochs):
            for i,(x, y) in enumerate(train_loader):
                optimizer.zero_grad()
                z = model(x.view(-1, 28 * 28))
                loss = criterion(z, y)
                loss.backward()
                optimizer.step()
                loss_accuracy['training_loss'].append(loss.data.item())
            correct = 0
            for x, y in validation_loader:
                yhat = model(x.view(-1, 28 * 28))
                _, label = torch.max(yhat, 1)
                correct += (label==y).sum().item()
            accuracy = 100 * (correct / len(validation_dataset))
            loss_accuracy['validation_accuracy'].append(accuracy)
        return loss_accuracy

    train_dataset = dsets.MNIST(root='./data', train=True, download=True, transform=transforms.ToTensor())
    validation_dataset = dsets.MNIST(root='./data', train=False, download=True, transform=transforms.ToTensor())
    train_loader = torch.utils.data.DataLoader(dataset=train_dataset, batch_size=2000, shuffle=True)
    validation_loader = torch.utils.data.DataLoader(dataset=validation_dataset, batch_size=5000, shuffle=False)
    criterion = nn.CrossEntropyLoss()
    input_dim = 28 * 28
    output_dim = 10
    layers = [input_dim, 100, 10, 100, 10, 100, output_dim]
    epochs = 15

    model = Net(layers)
    learning_rate = 0.01
    optimizer = torch.optim.SGD(model.parameters(), lr=learning_rate)
    training_results = train(model, criterion, train_loader, validation_loader, optimizer, epochs=epochs)

    model_Xavier = Net_Xavier(layers)
    optimizer = torch.optim.SGD(model_Xavier.parameters(), lr=learning_rate)
    training_results_Xavier = train(model_Xavier, criterion, train_loader, validation_loader, optimizer, epochs=epochs)

    model_Uniform = Net_Uniform(layers)
    optimizer = torch.optim.SGD(model_Uniform.parameters(), lr=learning_rate)
    training_results_Uniform = train(model_Uniform, criterion, train_loader, validation_loader, optimizer, epochs=epochs)

    plt.figure()
    plt.plot(training_results_Xavier['training_loss'], label='Xavier')
    plt.plot(training_results['training_loss'], label='Default')
    plt.plot(training_results_Uniform['training_loss'], label='Uniform')
    plt.ylabel('loss')
    plt.xlabel('iteration ')
    plt.title('training loss iterations')
    plt.legend()
    plt.savefig(f"{script_name}_loss.png")
    plt.close()

    plt.figure()
    plt.plot(training_results_Xavier['validation_accuracy'], label='Xavier')
    plt.plot(training_results['validation_accuracy'], label='Default')
    plt.plot(training_results_Uniform['validation_accuracy'], label='Uniform')
    plt.ylabel('validation accuracy')
    plt.xlabel('epochs')
    plt.legend()
    plt.savefig(f"{script_name}_valacc.png")
    plt.close()
    # --- End original script ---

output_text = f.getvalue()
doc.add_heading('Console Output', level=1)
doc.add_paragraph(output_text)

for img in [f"{script_name}_loss.png", f"{script_name}_valacc.png"]:
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(5.5))
        os.remove(img)

doc.save(word_output)
print(f"All output and plots saved to {word_output}") 