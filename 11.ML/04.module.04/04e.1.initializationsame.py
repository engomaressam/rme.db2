# <Initialization with Same Weights>
# Objective for this Notebook:
# 1. Learn how to Define the Neural Network with Same Weights Initialization, define Criterion Function, Optimizer, and Train the Model
# 2. Define the Neural Network with default Weights Initialization, define Criterion Function, Optimizer
# 3. Train the Model
# Table of Contents:
# - Neural Network Module and Training Function
# - Make Some Data
# - Define the Neural Network with Same Weights Initialization, define Criterion Function, Optimizer, and Train the Model
# Estimated Time Needed: 25 min
#
# Preparation
# We'll need the following libraries
import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import torch
import torch.nn as nn
from torch import sigmoid
import matplotlib.pyplot as plt
import numpy as np
torch.manual_seed(0)

script_name = os.path.splitext(os.path.basename(__file__))[0]
word_output = f"{script_name}.docx"
doc = Document()
doc.add_heading(f'Output for {script_name}', 0)

f = io.StringIO()
with redirect_stdout(f):
    # The function for plotting the model
    def PlotStuff(X, Y, model, epoch, leg=True):
        plt.plot(X.numpy(), model(X).detach().numpy(), label=('epoch ' + str(epoch)))
        plt.plot(X.numpy(), Y.numpy(), 'r')
        plt.xlabel('x')
        if leg == True:
            plt.legend()
        plt.savefig(f"{script_name}_plotstuff_{epoch}.png")
        plt.close()

    # Neural Network Module and Training Function
    class Net(nn.Module):
        def __init__(self, D_in, H, D_out):
            super(Net, self).__init__()
            self.linear1 = nn.Linear(D_in, H)
            self.linear2 = nn.Linear(H, D_out)
            self.a1 = None
            self.l1 = None
            self.l2 = None
        def forward(self, x):
            self.l1 = self.linear1(x)
            self.a1 = sigmoid(self.l1)
            self.l2 = self.linear2(self.a1)
            yhat = sigmoid(self.linear2(self.a1))
            return yhat

    def train(Y, X, model, optimizer, criterion, epochs=1000):
        cost = []
        for epoch in range(epochs):
            total = 0
            for y, x in zip(Y, X):
                yhat = model(x)
                loss = criterion(yhat, y)
                loss.backward()
                optimizer.step()
                optimizer.zero_grad()
                total += loss.item()
            cost.append(total)
            if epoch % 300 == 0:
                PlotStuff(X, Y, model, epoch, leg=True)
                model(X)
                plt.scatter(model.a1.detach().numpy()[:, 0], model.a1.detach().numpy()[:, 1], c=Y.numpy().reshape(-1))
                plt.title('activations')
                plt.savefig(f"{script_name}_activations_{epoch}.png")
                plt.close()
        return cost

    # Make Some Data
    X = torch.arange(-20, 20, 1).view(-1, 1).type(torch.FloatTensor)
    Y = torch.zeros(X.shape[0])
    Y[(X[:, 0] > -4) & (X[:, 0] < 4)] = 1.0

    # Define the loss function
    def criterion_cross(outputs, labels):
        out = -1 * torch.mean(labels * torch.log(outputs) + (1 - labels) * torch.log(1 - outputs))
        return out

    # Train the model with same weights initialization
    D_in = 1
    H = 2
    D_out = 1
    learning_rate = 0.1
    model = Net(D_in, H, D_out)

    # PyTorch default initialization
    model.state_dict()

    # Same Weights Initialization with all ones for weights and zeros for the bias.
    model.state_dict()['linear1.weight'][0] = 1.0
    model.state_dict()['linear1.weight'][1] = 1.0
    model.state_dict()['linear1.bias'][0] = 0.0
    model.state_dict()['linear1.bias'][1] = 0.0
    model.state_dict()['linear2.weight'][0] = 1.0
    model.state_dict()['linear2.bias'][0] = 0.0
    model.state_dict()

    # Optimizer, and Train the Model:
    optimizer = torch.optim.SGD(model.parameters(), lr=learning_rate)
    cost_cross = train(Y, X, model, optimizer, criterion_cross, epochs=1000)
    plt.plot(cost_cross)
    plt.xlabel('epoch')
    plt.title('cross entropy loss')
    plt.savefig(f"{script_name}_loss1.png")
    plt.close()

    # By examining the output of the parameters all though they have changed they are identical.
    model.state_dict()

    yhat = model(torch.tensor([[-2.0],[0.0],[2.0]]))
    yhat

    # Define the Neural Network, Criterion Function, Optimizer and Train the Model (default init)
    model = Net(D_in, H, D_out)
    optimizer = torch.optim.SGD(model.parameters(), lr=learning_rate)
    cost_cross = train(Y, X, model, optimizer, criterion_cross, epochs=1000)
    plt.plot(cost_cross)
    plt.xlabel('epoch')
    plt.title('cross entropy loss')
    plt.savefig(f"{script_name}_loss2.png")
    plt.close()

output_text = f.getvalue()
doc.add_heading('Console Output', level=1)
doc.add_paragraph(output_text)

# Collect all generated images
for img in sorted([f for f in os.listdir('.') if f.startswith(f"{script_name}_") and f.endswith('.png')]):
    doc.add_picture(img, width=Inches(5.5))
    os.remove(img)

doc.save(word_output)
print(f"All output and plots saved to {word_output}") 