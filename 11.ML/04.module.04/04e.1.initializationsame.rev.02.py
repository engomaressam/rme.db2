import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import torch
import torch.nn as nn
from torch import sigmoid
import matplotlib.pyplot as plt
import numpy as np

def PlotStuff(X, Y, model, epoch, leg=True):
    plt.plot(X.numpy(), model(X).detach().numpy(), label=('epoch ' + str(epoch)))
    plt.plot(X.numpy(), Y.numpy(), 'r')
    plt.xlabel('x')
    if leg == True:
        plt.legend()
    plt.savefig(f"{script_name}_plotstuff_{epoch}.png")
    plt.close()

class Net(nn.Module):
    def __init__(self, D_in, H, D_out):
        super(Net, self).__init__()
        self.linear1 = nn.Linear(D_in, H)
        self.linear2 = nn.Linear(H, D_out)
        self.a1 = None
        self.l1 = None
        self.l2 = None
    def forward(self, x):
        self.l1 = self.linear1(x)
        self.a1 = sigmoid(self.l1)
        self.l2 = self.linear2(self.a1)
        yhat = sigmoid(self.linear2(self.a1))
        return yhat

def train(Y, X, model, optimizer, criterion, epochs=1000):
    cost = []
    for epoch in range(epochs):
        total = 0
        for y, x in zip(Y, X):
            yhat = model(x)
            loss = criterion(yhat, y)
            loss.backward()
            optimizer.step()
            optimizer.zero_grad()
            total += loss.item()
        cost.append(total)
        if epoch % 300 == 0:
            PlotStuff(X, Y, model, epoch, leg=True)
            model(X)
            plt.scatter(model.a1.detach().numpy()[:, 0], model.a1.detach().numpy()[:, 1], c=Y.numpy().reshape(-1))
            plt.title('activations')
            plt.savefig(f"{script_name}_activations_{epoch}.png")
            plt.close()
    return cost

def criterion_cross(outputs, labels):
    out = -1 * torch.mean(labels * torch.log(outputs) + (1 - labels) * torch.log(1 - outputs))
    return out

script_name = os.path.splitext(os.path.basename(__file__))[0]
word_output = f"{script_name}.docx"
learning_rate = 0.05
X = torch.arange(-20, 20, 1).view(-1, 1).type(torch.FloatTensor)
Y = torch.zeros(X.shape[0])
Y[(X[:, 0] > -4) & (X[:, 0] < 4)] = 1.0
D_in = 1
H = 2
D_out = 1
model = Net(D_in, H, D_out)
optimizer = torch.optim.SGD(model.parameters(), lr=learning_rate)

doc = Document()
doc.add_heading(f'Output for {script_name}', 0)
f = io.StringIO()
with redirect_stdout(f):
    cost_cross = train(Y, X, model, optimizer, criterion_cross, epochs=1500)
    plt.plot(cost_cross)
    plt.xlabel('epoch')
    plt.title('cross entropy loss')
    plt.savefig(f"{script_name}_loss1.png")
    plt.close()
output_text = f.getvalue()
doc.add_heading('Console Output', level=1)
doc.add_paragraph(output_text)
for img in sorted([f for f in os.listdir('.') if f.startswith(f"{script_name}_") and f.endswith('.png')]):
    doc.add_picture(img, width=Inches(5.5))
    os.remove(img)
doc.save(word_output)
print(f"All output and plots saved to {word_output}") 