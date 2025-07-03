# <Using Dropout in Regression>
# Objective for this Notebook:
# 1. Create the Model and Cost Function the PyTorch way.
# 2. Learn Batch Gradient Descent
# Table of Contents:
# - Make Some Data
# - Create the Model and Cost Function the PyTorch way
# - Batch Gradient Descent
# Estimated Time Needed: 20 min
#
# Preparation
# We'll need the following libraries
import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import torch
import matplotlib.pyplot as plt
import torch.nn as nn
import torch.nn.functional as F
import numpy as np
from torch.utils.data import Dataset, DataLoader

script_name = os.path.splitext(os.path.basename(__file__))[0]
word_output = f"{script_name}.docx"
doc = Document()
doc.add_heading(f'Output for {script_name}', 0)

f = io.StringIO()
with redirect_stdout(f):
    torch.manual_seed(0)

    # Make Some Data
    # Create polynomial dataset class
    class Data(Dataset):
        def __init__(self, N_SAMPLES=40, noise_std=1, train=True):
            self.x = torch.linspace(-1, 1, N_SAMPLES).view(-1, 1)
            self.f = self.x ** 2
            if train != True:
                torch.manual_seed(1)
                self.y = self.f + noise_std * torch.randn(self.f.size())
                self.y = self.y.view(-1, 1)
                torch.manual_seed(0)
            else:
                self.y = self.f + noise_std * torch.randn(self.f.size())
                self.y = self.y.view(-1, 1)
        def __getitem__(self, index):
            return self.x[index], self.y[index]
        def __len__(self):
            return self.len
        def plot(self):
            plt.figure(figsize = (6.1, 10))
            plt.scatter(self.x.numpy(), self.y.numpy(), label="Samples")
            plt.plot(self.x.numpy(), self.f.numpy() ,label="True Function", color='orange')
            plt.xlabel("x")
            plt.ylabel("y")
            plt.xlim((-1, 1))
            plt.ylim((-2, 2.5))
            plt.legend(loc="best")
            plt.savefig(f"{script_name}_data.png")
            plt.close()

    data_set = Data()
    data_set.plot()

    # Create validation dataset object
    validation_set = Data(train=False)

    # Create the Model, Optimizer, and Total Loss Function (Cost)
    class Net(nn.Module):
        def __init__(self, in_size, n_hidden, out_size, p=0):
            super(Net, self).__init__()
            self.drop = nn.Dropout(p=p)
            self.linear1 = nn.Linear(in_size, n_hidden)
            self.linear2 = nn.Linear(n_hidden, n_hidden)
            self.linear3 = nn.Linear(n_hidden, out_size)
        def forward(self, x):
            x = F.relu(self.drop(self.linear1(x)))
            x = F.relu(self.drop(self.linear2(x)))
            x = self.linear3(x)
            return x

    # Create the model objects
    model = Net(1, 300, 1)
    model_drop = Net(1, 300, 1, p=0.5)

    # Set the optimizer and criterion function
    optimizer_ofit = torch.optim.Adam(model.parameters(), lr=0.01)
    optimizer_drop = torch.optim.Adam(model_drop.parameters(), lr=0.01)
    criterion = torch.nn.MSELoss()

    # Initialize the dict to contain the loss results
    LOSS={}
    LOSS['training data no dropout']=[]
    LOSS['validation data no dropout']=[]
    LOSS['training data dropout']=[]
    LOSS['validation data dropout']=[]

    epochs = 500

    def train_model(epochs):
        for epoch in range(epochs):
            yhat = model(data_set.x)
            yhat_drop = model_drop(data_set.x)
            loss = criterion(yhat, data_set.y)
            loss_drop = criterion(yhat_drop, data_set.y)
            LOSS['training data no dropout'].append(loss.item())
            LOSS['validation data no dropout'].append(criterion(model(validation_set.x), validation_set.y).item())
            LOSS['training data dropout'].append(loss_drop.item())
            model_drop.eval()
            LOSS['validation data dropout'].append(criterion(model_drop(validation_set.x), validation_set.y).item())
            model_drop.train()
            optimizer_ofit.zero_grad()
            optimizer_drop.zero_grad()
            loss.backward()
            loss_drop.backward()
            optimizer_ofit.step()
            optimizer_drop.step()
    train_model(epochs)

    # Set the model with dropout to evaluation mode
    model_drop.eval()

    # Make the prediction
    yhat = model(data_set.x)
    yhat_drop = model_drop(data_set.x)

    # Plot the predictions for both models
    plt.figure(figsize=(6.1, 10))
    plt.scatter(data_set.x.numpy(), data_set.y.numpy(), label="Samples")
    plt.plot(data_set.x.numpy(), data_set.f.numpy(), label="True function", color='orange')
    plt.plot(data_set.x.numpy(), yhat.detach().numpy(), label='no dropout', c='r')
    plt.plot(data_set.x.numpy(), yhat_drop.detach().numpy(), label="dropout", c ='g')
    plt.xlabel("x")
    plt.ylabel("y")
    plt.xlim((-1, 1))
    plt.ylim((-2, 2.5))
    plt.legend(loc = "best")
    plt.savefig(f"{script_name}_preds.png")
    plt.close()

    # Plot the loss
    plt.figure(figsize=(6.1, 10))
    for key, value in LOSS.items():
        plt.plot(np.log(np.array(value)), label=key)
        plt.legend()
        plt.xlabel("iterations")
        plt.ylabel("Log of cost or total loss")
    plt.savefig(f"{script_name}_loss.png")
    plt.close()

output_text = f.getvalue()
doc.add_heading('Console Output', level=1)
doc.add_paragraph(output_text)

for img in [f"{script_name}_data.png", f"{script_name}_preds.png", f"{script_name}_loss.png"]:
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(5.5))
        os.remove(img)

doc.save(word_output)
print(f"All output and plots saved to {word_output}")

# End of script 