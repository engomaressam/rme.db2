import sys
import io
import os
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Inches
import torch
import matplotlib.pyplot as plt
import torch.nn as nn
import torch.nn.functional as F
import numpy as np
from matplotlib.colors import ListedColormap
from torch.utils.data import Dataset, DataLoader

# The function for plotting the diagram

def plot_decision_regions_3class(data_set, model=None):
    cmap_light = ListedColormap([ '#0000FF','#FF0000'])
    cmap_bold = ListedColormap(['#FF0000', '#00FF00', '#00AAFF'])
    X = data_set.x.numpy()
    y = data_set.y.numpy()
    h = .02
    x_min, x_max = X[:, 0].min() - 0.1, X[:, 0].max() + 0.1
    y_min, y_max = X[:, 1].min() - 0.1, X[:, 1].max() + 0.1
    xx, yy = np.meshgrid(np.arange(x_min, x_max, h), np.arange(y_min, y_max, h))
    newdata = np.c_[xx.ravel(), yy.ravel()]
    Z = data_set.multi_dim_poly(newdata).flatten()
    f = np.zeros(Z.shape)
    f[Z > 0] = 1
    f = f.reshape(xx.shape)
    if model != None:
        model.eval()
        XX = torch.Tensor(newdata)
        _, yhat = torch.max(model(XX), 1)
        yhat = yhat.numpy().reshape(xx.shape)
        plt.pcolormesh(xx, yy, yhat, cmap=cmap_light)
        plt.contour(xx, yy, f, cmap=plt.cm.Paired)
    else:
        plt.contour(xx, yy, f, cmap=plt.cm.Paired)
        plt.pcolormesh(xx, yy, f, cmap=cmap_light)
    plt.title("decision region vs True decision boundary")

def accuracy(model, data_set):
    _, yhat = torch.max(model(data_set.x), 1)
    return (yhat == data_set.y).numpy().mean()

# Data class for creating dataset object
class Data(Dataset):
    def __init__(self, N_SAMPLES=1000, noise_std=0.15, train=True):
        a = np.matrix([-1, 1, 2, 1, 1, -3, 1]).T
        self.x = np.matrix(np.random.rand(N_SAMPLES, 2))
        self.f = np.array(a[0] + (self.x) * a[1:3] + np.multiply(self.x[:, 0], self.x[:, 1]) * a[4] + np.multiply(self.x, self.x) * a[5:7]).flatten()
        self.a = a
        self.y = np.zeros(N_SAMPLES)
        self.y[self.f > 0] = 1
        self.y = torch.from_numpy(self.y).type(torch.LongTensor)
        self.x = torch.from_numpy(self.x).type(torch.FloatTensor)
        self.x = self.x + noise_std * torch.randn(self.x.size())
        self.f = torch.from_numpy(self.f)
        self.a = a
        if train == True:
            torch.manual_seed(1)
            self.x = self.x + noise_std * torch.randn(self.x.size())
            torch.manual_seed(0)
        self.len = N_SAMPLES
    def __getitem__(self, index):
        return self.x[index], self.y[index]
    def __len__(self):
        return self.len
    def plot(self):
        X = self.x.numpy()
        y = self.y.numpy()
        h = .02
        x_min, x_max = X[:, 0].min(), X[:, 0].max()
        y_min, y_max = X[:, 1].min(), X[:, 1].max()
        xx, yy = np.meshgrid(np.arange(x_min, x_max, h), np.arange(y_min, y_max, h))
        Z = self.multi_dim_poly(np.c_[xx.ravel(), yy.ravel()]).flatten()
        f = np.zeros(Z.shape)
        f[Z > 0] = 1
        f = f.reshape(xx.shape)
        plt.title('True decision boundary  and sample points with noise ')
        plt.plot(self.x[self.y == 0, 0].numpy(), self.x[self.y == 0,1].numpy(), 'bo', label='y=0')
        plt.plot(self.x[self.y == 1, 0].numpy(), self.x[self.y == 1,1].numpy(), 'ro', label='y=1')
        plt.contour(xx, yy, f,cmap=plt.cm.Paired)
        plt.xlim(0,1)
        plt.ylim(0,1)
        plt.legend()
    def multi_dim_poly(self, x):
        x = np.matrix(x)
        out = np.array(self.a[0] + (x) * self.a[1:3] + np.multiply(x[:, 0], x[:, 1]) * self.a[4] + np.multiply(x, x) * self.a[5:7])
        out = np.array(out)
        return out

# Net class
class Net(nn.Module):
    def __init__(self, in_size, n_hidden, out_size, p=0):
        super(Net, self).__init__()
        self.drop = nn.Dropout(p=p)
        self.linear1 = nn.Linear(in_size, n_hidden)
        self.linear2 = nn.Linear(n_hidden, n_hidden)
        self.linear3 = nn.Linear(n_hidden, out_size)
    def forward(self, x):
        x = F.relu(self.drop(self.linear1(x)))
        x = F.relu(self.drop(self.linear2(x)))
        x = self.linear3(x)
        return x

# --- Main script logic below ---
script_name = os.path.splitext(os.path.basename(__file__))[0]
word_output = f"{script_name}.docx"
doc = Document()
doc.add_heading(f'Output for {script_name}', 0)
f = io.StringIO()
with redirect_stdout(f):
    data_set = Data(noise_std=0.2)
    data_set.plot()
    plt.savefig(f"{script_name}_data.png")
    plt.close()
    torch.manual_seed(0)
    validation_set = Data(train=False)
    model_drop = Net(2, 300, 2, p=0.7)  # changed dropout from 0.5 to 0.7
    optimizer_drop = torch.optim.Adam(model_drop.parameters(), lr=0.01)
    criterion = torch.nn.CrossEntropyLoss()
    LOSS = []
    epochs = 700  # changed from 500
    for epoch in range(epochs):
        yhat_drop = model_drop(data_set.x)
        loss_drop = criterion(yhat_drop, data_set.y)
        LOSS.append(loss_drop.item())
        optimizer_drop.zero_grad()
        loss_drop.backward()
        optimizer_drop.step()
    model_drop.eval()
    print("The accuracy of the model with dropout (p=0.7): ", accuracy(model_drop, validation_set))
    plot_decision_regions_3class(data_set, model_drop)
    plt.savefig(f"{script_name}_decision.png")
    plt.close()
    plt.figure(figsize=(6.1, 10))
    plt.plot(np.log(np.array(LOSS)), label='training data dropout')
    plt.legend()
    plt.xlabel("iterations")
    plt.ylabel("Log of cost or total loss")
    plt.savefig(f"{script_name}_loss.png")
    plt.close()
output_text = f.getvalue()
doc.add_heading('Console Output', level=1)
doc.add_paragraph(output_text)
for img in [f"{script_name}_loss.png", f"{script_name}_data.png", f"{script_name}_decision.png"]:
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(5.5))
        os.remove(img)
doc.save(word_output)
print(f"All output and plots saved to {word_output}") 